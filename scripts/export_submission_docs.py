from pathlib import Path
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / 'docs' / 'final_submission_report.md'
DOCX_OUT = ROOT / 'docs' / 'final_submission_report.docx'
PDF_OUT = ROOT / 'docs' / 'final_submission_report.pdf'

content = SOURCE.read_text(encoding='utf-8')

# Generate DOCX
word_doc = Document()
for line in content.splitlines():
    if line.startswith('# '):
        word_doc.add_heading(line[2:], level=1)
    elif line.startswith('## '):
        word_doc.add_heading(line[3:], level=2)
    elif line.strip():
        word_doc.add_paragraph(line)
word_doc.save(DOCX_OUT)

# Generate PDF
pdf_canvas = canvas.Canvas(str(PDF_OUT), pagesize=letter)
text = pdf_canvas.beginText(40, 750)
for line in content.splitlines():
    if line.startswith('# '):
        text.setFont('Helvetica-Bold', 16)
        text.textLine(line[2:])
    elif line.startswith('## '):
        text.setFont('Helvetica-Bold', 12)
        text.textLine(line[3:])
    elif line.strip():
        text.setFont('Helvetica', 11)
        text.textLine(line)
pdf_canvas.drawText(text)
pdf_canvas.save()

print(f'Created {DOCX_OUT.name} and {PDF_OUT.name}')
